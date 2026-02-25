import os

with open("build_report.py", "r", encoding="utf-8") as f:
    code = f.read()

# 1. Add footer helpers and imports
import_insert = """from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
"""
if "from docx.oxml.shared import OxmlElement" not in code:
    code = code.replace("from docx.shared import Pt, Inches, RGBColor", "from docx.shared import Pt, Inches, RGBColor\n" + import_insert)

footer_func = """
def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run("School of CSA, REVA University, Bengaluru")
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(128, 128, 128)
        
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '4')
        top.set(qn('w:space'), '1')
        top.set(qn('w:color'), 'auto')
        pBdr.append(top)
        pPr.append(pBdr)

"""
if "def add_footer(doc):" not in code:
    code = code.replace("def main_doc:", footer_func + "def main_doc:")
    # or just before if __name__ == "__main__":
    code = code.replace('if __name__ == "__main__":', footer_func + 'if __name__ == "__main__":')

# 2. Add add_footer(main_doc) before save
if "add_footer(main_doc)" not in code:
    code = code.replace('main_doc.save("segT_Major_Project_Report.docx")', 'add_footer(main_doc)\n    main_doc.save("segT_Major_Project_Report.docx")')

# 3. Replace 7.5 and 7.6
old_chapter_7_UI = """    add_sub_heading(doc, "7.5 USER INTERFACE AND VISUALIZATION")
    add_paragraph(doc, "The graphical user interface (GUI) of the 'Terrain Safety Analysis with SegFormer' application is built using the Gradio web framework. It provides real-time, interactive terrain analysis for autonomous navigation. The interface follows a modern, dark-themed dashboard layout and is divided into two main sections:")
    add_paragraph(doc, "1. The Navigation & Control Sidebar (Left Panel): Serves as the primary interaction hub. It includes an Input Image module for uploading raw RGB terrain captures, Settings for adjusting HUD opacity, and Advanced Configuration Menus for dynamic JSON class mapping and turning features like Depth and Pathfinding on or off.")
    add_paragraph(doc, "2. The Output Visualization Grid (Right Panels): Structured as a 2x2 grid offering distinct perspectives. The 'HUD Prediction' panel overlays the binary safety mask (Red for Hazards, Green for Safe) and optimal traversal path over the original image. The 'Raw Mask' panel displays the 150-class semantic segmentation mask in greyscale. The 'Depth Map' displays the monocular depth estimation in a thermal pseudocolour format.")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("\\n\\n[ INSERT GRADIO UI APP SCREENSHOT HERE ]\\n\\n")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.italic = True
    add_table_title(doc, "Figure 7.1: Interactive Gradio Web Dashboard (Navigation & 2D Outputs)")

    add_sub_heading(doc, "7.6 3D TERRAIN MESH AND TELEMETRY")
    add_paragraph(doc, "Beyond 2D projections, the dashboard provides deep geometric context and precise telemetry outputs.")
    add_paragraph(doc, "1. 3D Terrain View (Sci-Fi Mode): This module acts as the ultimate synthesis of semantic and geometric data. It extrudes the 2D RGB image into a 3D topographic model using metric values from Depth Anything V2. The semantic safety mask is draped over this geometry as a wireframe (red for elevated hazards, green for safe valleys), along with a 3D pathfinding trajectory.")
    add_paragraph(doc, "2. High-Level Telemetry and Detailed Stats: A prominent banner displays the immediate 'Safety Score' (e.g., 35.9%), indicating the percentage of visible terrain deemed traversable. The 'Detailed Stats' terminal outputs precise JSON telemetry mirroring W&B logic, containing exact pixel analytics (safe vs hazard pixel counts distributions), and Mean Confidence metrics for the SegFormer inference predictions.")

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("\\n\\n[ INSERT 3D STATS UI SCREENSHOT HERE ]\\n\\n")
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)
    r2.italic = True
    add_table_title(doc, "Figure 7.2: 3D Terrain Mesh & Detailed Inference Telemetry")"""

new_chapter_7_UI = """    add_sub_heading(doc, "7.5 USER INTERFACE AND VISUALIZATION")
    add_paragraph(doc, "The graphical user interface (GUI) of the 'Terrain Safety Analysis with SegFormer (v2)' application is built using the Gradio web framework. It provides real-time, interactive terrain analysis for autonomous navigation. The primary objective of this system is to ingest a standard 2D RGB terrain image and process it through a deep learning pipeline (SegFormer for semantic segmentation and Depth Anything for metric depth estimation). The resulting output categorizes the environment into traversable 'Safe' zones versus impassable 'Hazard' regions, visualizing these boundaries directly to the operator.")
    add_paragraph(doc, "The interface follows a modern, dark-themed, dashboard-style layout, optimizing screen real estate to present multiple streams of visual data concurrently. It is logically divided into two main sections:")
    add_paragraph(doc, "1. The Navigation & Control Sidebar (Left panel)")
    add_paragraph(doc, "2. The Output Visualization Grid (Right and Center panels)")

    add_sub_heading(doc, "7.5.1 The Navigation & Control Sidebar")
    add_paragraph(doc, "The left sidebar serves as the primary interaction hub for the human operator. It is divided into several collapsible accordion menus and interactive controls, allowing for dynamic parameter tuning without requiring the user to write or modify backend code.")
    
    add_paragraph(doc, "Input Image Module:")
    add_paragraph(doc, "At the top left, the 'Input Image' window allows the user to upload or drag-and-drop the raw RGB image intended for analysis. Below the image, standard Gradio controls are visible for uploading, capturing from a webcam, or clearing the current image.")
    
    add_paragraph(doc, "Settings and Overlays:")
    add_paragraph(doc, "Directly beneath the input module is the 'Settings' tab. This section displays a slider for HUD Opacity. This slider controls the transparency of the semantic overlays rendered in the output panels. By allowing the operator to adjust opacity, the system ensures that the underlying raw image features (like hidden rocks or textures) can still be inspected beneath the colored masks.")
    
    add_paragraph(doc, "Advanced Configuration Menus:")
    add_paragraph(doc, "• Class Mapping (JSON): Allows the operator to dynamically override the semantic mapping dictionary. For example, a user could re-classify the 'sand' class from a 'Safe' zone to a 'Hazard' zone depending on the rover's drivetrain capabilities.")
    add_paragraph(doc, "• Advanced Features: Displays checkboxes to Enable Depth pipeline, Enable Pathfinding (A* or Dijkstra), and Show 3D View.")
    add_paragraph(doc, "• Rover Settings: Houses parameters for the physical dimensions of the rover (e.g., chassis width). This is crucial for the 'Rover Dimension Constraint' algorithm, ensuring paths are physically wide enough to traverse.")
    add_paragraph(doc, "At the very bottom of the sidebar is a prominent orange 'Analyze Terrain' button, which acts as the execution trigger for the ML pipeline once all parameters are set.")

    add_sub_heading(doc, "7.5.2 The Output Visualization Grid")
    add_paragraph(doc, "The right side of the dashboard is dedicated to displaying the processed outputs from the deep learning models. It is structured into a 2x2 style grid, offering distinct visual perspectives.")
    
    add_paragraph(doc, "HUD Prediction (with Path) - Top Left:")
    add_paragraph(doc, "This panel overlays the binary safety mask directly onto the original rocky terrain image. The Red Zones (Hazard) demarcate areas deemed impassable. The Green Zones (Safe) indicate navigable terrain. A distinct dotted blue pathfinding line traces an optimal route through the safe zone.")
    
    add_paragraph(doc, "Raw Mask - Top Right:")
    add_paragraph(doc, "This panel displays the raw semantic segmentation mask generated directly by the SegFormer model in greyscale. The varying shades of grey represent the different semantic classes (out of the 150 ADE20K classes) the model has identified, demonstrating its ability to separate textural entities prior to safety classification.")
    
    add_paragraph(doc, "Depth Map - Bottom Center:")
    add_paragraph(doc, "This large module displays the output of the monocular depth estimation model using a Thermal Pseudocolour Representation. Red and orange hues denote close physical proximity, while blue and purple hues denote distant background horizons. The optimal traversal path is superimposed entirely in bright cyan across this depth map.")
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("\\n\\n[ INSERT GRADIO UI APP SCREENSHOT HERE ]\\n\\n")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.italic = True
    add_table_title(doc, "Figure 7.1: Interactive Gradio Web Dashboard (Navigation & 2D Outputs)")

    add_sub_heading(doc, "7.6 3D TERRAIN MESH AND TELEMETRY")
    add_paragraph(doc, "While the primary viewport focuses on 2D semantic overlays and HUDs, the downward-scrolled section provides deep geometric context and precise, quantifiable telemetry derived from the inference pipeline.")
    
    add_paragraph(doc, "3D Terrain View (Sci-Fi Mode):")
    add_paragraph(doc, "At the top of the interface is the 3D Terrain Mesh viewer. The system has taken the 2D RGB image and extruded it into a three-dimensional topographic model using the metric values from Depth Anything V2. The 2D semantic safety mask is draped directly over this 3D geometry as a wireframe. Elevational hazards form a red wireframe, while the flatter traversable valleys form a green wireframe. A vividly plotted segmented line traces the 3D pathfinding route.")
    
    add_paragraph(doc, "High-Level Telemetry:")
    add_paragraph(doc, "Directly beneath the 3D viewer is a prominent banner displaying the 'Safety Score' (e.g., 35.9%). This metric provides the operator with an immediate understanding of the overall navigability of the current scene, indicating the percentage of visible terrain deemed traversable.")
    
    add_paragraph(doc, "Detailed Stats (JSON Output):")
    add_paragraph(doc, "A raw data terminal outputs the precise numerical telemetry generated during the inference cycle. The JSON data dictionary contains several critical metrics, including pixel analytics (safe vs hazard counts), distribution percentages, and confidence metrics (mean confidence of the SegFormer model across predictions).")

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("\\n\\n[ INSERT 3D STATS UI SCREENSHOT HERE ]\\n\\n")
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)
    r2.italic = True
    add_table_title(doc, "Figure 7.2: 3D Terrain Mesh & Detailed Inference Telemetry")"""

code = code.replace(old_chapter_7_UI, new_chapter_7_UI)


# 4. Expand Chapter 9
old_chapter_9_intro = """    add_sub_heading(doc, "9.1 SUMMARY OF CONTRIBUTIONS")"""
new_chapter_9_intro = """    add_paragraph(doc, "The rapid proliferation of autonomous systems over the past decade has fundamentally redefined the landscape of modern robotics. Historically, establishing environmental awareness for robotic navigation relied heavily on active sensory hardware, predominantly Light Detection and Ranging (LiDAR) and sophisticated stereo-camera rigs. While these technologies offer high-fidelity spatial awareness, they introduce substantial barriers in terms of financial cost, mechanical complexity, and power consumption. These limitations have largely confined advanced autonomous navigation to well-funded research laboratories or high-budget industrial applications, severely hindering the democratization of intelligent, self-navigating robotics in resource-constrained sectors such as agriculture, search and rescue operations, and low-cost exploratory missions.")
    add_paragraph(doc, "To address this critical bottleneck, this project proposed and developed a radically alternative paradigm: achieving robust spatial and semantic awareness using only a single, standard monocular RGB camera. By discarding active sensors in favor of advanced deep learning architectures, the system shifts the computational burden from mechanical hardware to intelligent software. The emergence of Vision Transformers (ViTs) and self-supervised depth estimation models has made this transition viable. Transformers, originally designed for natural language processing, possess an inherent ability to capture global contextual information across an image, decisively outperforming traditional Convolutional Neural Networks (CNNs) in unstructured, natural environments where rigid geometric boundaries do not exist.")
    add_paragraph(doc, "This report has detailed the end-to-end development of the 'Terrain Safety Analysis' system. The core of this system is an innovative algorithmic bridge that fuses the abstract, multi-dimensional semantic logits outputted by a pre-trained SegFormer model with the concrete, metric depth gradients produced by the Depth Anything V2 model. This fusion successfully grounds pixel-level semantic classifications into physical, real-world geometry. The resulting data is ingested by a custom Safety Logic Engine, strictly governed by a configurable JSON schema, which aggressively evaluates transversability against the physical dimensional constraints of the host rover.")
    add_paragraph(doc, "Furthermore, this project placed a heavy emphasis on Human-Computer Interaction (HCI) and MLOps observability. The complex PyTorch backend is entirely encapsulated within an intuitive Gradio web dashboard, allowing operators to visually interpret 2D semantic HUDs, explore 3D topological meshes, and dynamically adjust safety classifications without requiring any software engineering expertise. The integration of continuous real-time telemetry streaming to Weights & Biases ensures that the deployed models can be audited and refined over time.")
    
    add_sub_heading(doc, "9.1 SUMMARY OF CONTRIBUTIONS")"""

code = code.replace(old_chapter_9_intro, new_chapter_9_intro)


# 5. Expand Chapter 10
old_chapter_10 = """    add_paragraph(doc, "[11] T. Abid et al. (2021). 'Gradio: Hassle-Free Sharing and Testing of ML Models in the Wild'.")"""

new_chapter_10 = """    add_paragraph(doc, "[11] T. Abid et al. (2021). 'Gradio: Hassle-Free Sharing and Testing of ML Models in the Wild'.")

    add_sub_heading(doc, "Online References and Documentation")
    add_paragraph(doc, "[12] Hugging Face. (2024). Transformers Documentation. Retrieved from https://huggingface.co/docs/transformers")
    add_paragraph(doc, "[13] PyTorch Contributors. (2024). PyTorch Documentation. Retrieved from https://pytorch.org/docs/stable/index.html")
    add_paragraph(doc, "[14] Gradio Team. (2024). Gradio Unifying Framework Documentation. Retrieved from https://www.gradio.app/docs/")
    add_paragraph(doc, "[15] Weights & Biases. (2024). Experiment Tracking API and MLOps Guide. Retrieved from https://docs.wandb.ai/")
    add_paragraph(doc, "[16] NVIDIA. (2024). TensorRT SDK Documentation for Deep Learning Inference. Retrieved from https://developer.nvidia.com/tensorrt")
    add_paragraph(doc, "[17] ONNX Runtime. (2024). Cross-Platform Machine Learning Inferencing. Retrieved from https://onnxruntime.ai/docs/")
    add_paragraph(doc, "[18] MIT CSAIL. (2024). ADE20K Dataset Outline and Benchmarks. Retrieved from https://groups.csail.mit.edu/vision/datasets/ADE20K/")
    add_paragraph(doc, "[19] Python Software Foundation. (2024). The Python Language Reference. Retrieved from https://docs.python.org/3/reference/")
    add_paragraph(doc, "[20] OpenCV Team. (2024). OpenCV-Python Tutorials and API. Retrieved from https://docs.opencv.org/4.x/")
    add_paragraph(doc, "[21] SciPy Community. (2024). NumPy Reference Guide. Retrieved from https://numpy.org/doc/stable/")
    add_paragraph(doc, "[22] Matplotlib Development Team. (2024). Matplotlib: Visualization with Python. Retrieved from https://matplotlib.org/")
    add_paragraph(doc, "[23] Scikit-Learn Developers. (2024). Scikit-learn: Machine Learning in Python. Retrieved from https://scikit-learn.org/stable/")"""

code = code.replace(old_chapter_10, new_chapter_10)

with open("build_report.py", "w", encoding="utf-8") as f:
    f.write(code)
    
print("Patch executed successfully!")

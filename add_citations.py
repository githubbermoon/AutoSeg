"""
add_citations.py
Adds inline citation markers [n] to the body text of segT_Major_Project_Report.docx.
Only adds citations — does NOT modify any other text, styling, or formatting.
Saves output to a NEW file to preserve the original.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
import copy, re

INPUT_FILE = "segT_Major_Project_Report.docx"
OUTPUT_FILE = "segT_Major_Project_Report_cited.docx"

# ==============================
# CITATION MAP: keyword → [ref_number]
# Each keyword is mapped to the bibliography entry it should cite.
# Only the first occurrence per paragraph gets the citation.
# ==============================
CITATION_MAP = {
    # [1] FCN - Long et al., 2015
    "FCN": "[1]",
    "Fully Convolutional Network": "[1]",
    
    # [2] U-Net - Ronneberger et al., 2015
    "U-Net": "[2]",
    
    # [3] ViT - Dosovitskiy et al., 2020
    "Vision Transformer": "[3]",
    "ViT": "[3]",
    
    # [4] SegFormer - Xie et al., 2021
    "SegFormer": "[4]",
    
    # [5] Depth Anything - Yang et al., 2024
    "Depth Anything": "[5]",
    
    # [6] ADE20K - Zhou et al., 2017 (academic)
    "scene parsing": "[6]",
    
    # [7] ResNet - He et al., 2016
    "ResNet": "[7]",
    "residual learning": "[7]",
    "deep residual": "[7]",
    "Convolutional Neural Network": "[7]",
    
    # [8] Attention / Transformer - Vaswani et al., 2017
    "Attention is all you need": "[8]",
    "self-attention mechanism": "[8]",
    
    # [9] DeepLab - Chen et al., 2018
    "DeepLab": "[9]",
    "DeepLabV3": "[9]",
    "DeepLabv3": "[9]",

    # [10] SegNet - Badrinarayanan et al., 2017
    "SegNet": "[10]",
    "encoder-decoder architecture": "[10]",
    "encoder-decoder": "[10]",
    
    # [11] Gradio paper - Abid et al., 2021 (academic)
    "Gradio: Hassle-Free": "[11]",
    "Gradio library": "[11]",
    
    # [12] Hugging Face
    "Hugging Face": "[12]",
    "HuggingFace": "[12]",
    "pre-trained model": "[12]",
    
    # [13] PyTorch
    "PyTorch": "[13]",
    
    # [14] Gradio docs
    "Gradio": "[14]",
    
    # [15] W&B
    "Weights & Biases": "[15]",
    "W&B": "[15]",
    
    # [16] TensorRT
    "TensorRT": "[16]",
    
    # [17] ONNX
    "ONNX": "[17]",
    
    # [18] ADE20K MIT CSAIL
    "ADE20K": "[18]",
    "ADE20k": "[18]",
    
    # [19] Python
    "Python 3": "[19]",
    "Python": "[19]",
    
    # [20] OpenCV
    "OpenCV": "[20]",
    "opencv": "[20]",
    
    # [21] NumPy
    "NumPy": "[21]",
    "numpy": "[21]",
    "numerical computation": "[21]",
    "numerical computations": "[21]",
    "numerical telemetry": "[21]",
    
    # [22] Matplotlib / Plotly
    "Matplotlib": "[22]",
    "Plotly": "[22]",
    
    # [23] Scikit
    "scikit": "[23]",
    "skimage": "[23]",
    "graph-based": "[23]",
    "route_through_array": "[23]",
    "traversal path": "[23]",
    "A* traversal": "[23]",
}

# Sections to SKIP (don't add citations to these)
SKIP_SECTIONS = [
    "BIBLIOGRAPHY",
    "APPENDIX",
    "Sample Source Code",
    "Online References",
    "CHAPTER 10",
    "CHAPTER 11",
]


def already_has_citation(text, citation):
    """Check if citation already exists near the keyword."""
    return citation in text


def add_citation_to_run(paragraph, keyword, citation):
    """
    Finds 'keyword' in the paragraph's runs and inserts 'citation' 
    as bold blue text right after it. Only modifies the first occurrence.
    Returns True if citation was added.
    """
    full_text = paragraph.text
    
    # Skip if keyword not in text or citation already present
    if keyword not in full_text:
        return False
    if already_has_citation(full_text, citation):
        return False
    
    runs = paragraph.runs
    if not runs:
        return False
    
    # Concatenate run texts to find position
    concat = ""
    run_boundaries = []
    for r in runs:
        start = len(concat)
        concat += r.text
        run_boundaries.append((start, len(concat), r))
    
    # Find keyword position in concatenated text
    pos = concat.find(keyword)
    if pos == -1:
        return False
    
    end_pos = pos + len(keyword)
    
    # Find which run the keyword ends in
    target_run = None
    insert_pos_in_run = None
    for start, end, run in run_boundaries:
        if end_pos <= end and end_pos > start:
            target_run = run
            insert_pos_in_run = end_pos - start
            break
    
    if target_run is None:
        return False
    
    # Split the target run's text at the insert position
    original_text = target_run.text
    before = original_text[:insert_pos_in_run]
    after = original_text[insert_pos_in_run:]
    
    # Set the run's text to the 'before' part
    target_run.text = before
    
    # Create a NEW run for the citation
    citation_run = copy.deepcopy(target_run._element)
    citation_run.text = " " + citation
    
    # Style the citation: bold blue, same size as body text
    rpr = citation_run.find(qn('w:rPr'))
    if rpr is None:
        rpr = citation_run.makeelement(qn('w:rPr'), {})
        citation_run.insert(0, rpr)
    
    # Remove any superscript vertAlign (keep it inline)
    for existing in rpr.findall(qn('w:vertAlign')):
        rpr.remove(existing)
    
    # Set color to blue
    color_elem = rpr.find(qn('w:color'))
    if color_elem is None:
        color_elem = rpr.makeelement(qn('w:color'), {})
        rpr.append(color_elem)
    color_elem.set(qn('w:val'), '0000FF')
    
    # Make bold
    bold_elem = rpr.find(qn('w:b'))
    if bold_elem is None:
        bold_elem = rpr.makeelement(qn('w:b'), {})
        rpr.append(bold_elem)
    
    # Remove any font size override — use same size as surrounding text
    for tag in [qn('w:sz'), qn('w:szCs')]:
        elem = rpr.find(tag)
        if elem is not None:
            rpr.remove(elem)
    
    # Create a run for the 'after' text (preserve original formatting)
    if after:
        after_run = copy.deepcopy(target_run._element)
        after_run.text = after
        after_rpr = after_run.find(qn('w:rPr'))
        if after_rpr is not None:
            for v in after_rpr.findall(qn('w:vertAlign')):
                after_rpr.remove(v)
    
    # Insert the elements in order
    parent = target_run._element.getparent()
    idx = list(parent).index(target_run._element)
    parent.insert(idx + 1, citation_run)
    if after:
        parent.insert(idx + 2, after_run)
    
    return True


def main():
    doc = Document(INPUT_FILE)
    
    citations_added = []
    in_skip_section = False
    
    # ============================
    # PASS 1: Body paragraphs
    # ============================
    for para in doc.paragraphs:
        text = para.text.strip()
        
        if any(skip in text for skip in SKIP_SECTIONS):
            in_skip_section = True
        
        if in_skip_section:
            continue
        if not text or len(text) < 20:
            continue
        if text.startswith(('import ', 'from ', 'def ', '#', '    ')):
            continue
        
        for keyword, citation in CITATION_MAP.items():
            if keyword in text:
                success = add_citation_to_run(para, keyword, citation)
                if success:
                    citations_added.append((citation, keyword, text[:60]))
    
    # ============================
    # PASS 2: Table cells (Software Requirements, etc.)
    # ============================
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if not text or len(text) < 5:
                        continue
                    for keyword, citation in CITATION_MAP.items():
                        if keyword in text:
                            success = add_citation_to_run(para, keyword, citation)
                            if success:
                                citations_added.append((citation, keyword, f"[TABLE] {text[:50]}"))
    
    # Save to new file
    doc.save(OUTPUT_FILE)
    
    print(f"✅ Saved cited report to: {OUTPUT_FILE}")
    print(f"\n📌 Citations added: {len(citations_added)}")
    for cit, kw, ctx in citations_added:
        print(f"  {cit} for '{kw}' → ...{ctx}...")


if __name__ == "__main__":
    main()
